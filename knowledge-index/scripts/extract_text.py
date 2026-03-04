#!/usr/bin/env python3
"""
本地文档文本提取脚本

用法:
    python extract_text.py <file_path> [--format json|text]
    python extract_text.py document.pdf --format json
    python extract_text.py document.docx

输出 (JSON 格式):
    {
        "text": "提取的纯文本内容",
        "metadata": {
            "page_count": 10,
            "char_count": 5000,
            "extractor": "pymupdf"
        },
        "success": true
    }
"""

import sys
import json
import os
from pathlib import Path


def extract_pdf_pymupdf(file_path: str) -> dict:
    """使用 PyMuPDF 提取 PDF 文本"""
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(file_path)
        text_parts = []
        page_count = len(doc)

        for page_num, page in enumerate(doc):
            text = page.get_text()
            if text.strip():
                text_parts.append(f"[第{page_num + 1}页]\n{text}")

        doc.close()

        return {
            "text": "\n\n".join(text_parts),
            "metadata": {
                "page_count": page_count,
                "extractor": "pymupdf",
                "char_count": sum(len(t) for t in text_parts)
            },
            "success": True
        }
    except ImportError:
        return {"success": False, "error": "PyMuPDF 未安装，请运行: pip install PyMuPDF"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def extract_pdf_pdfplumber(file_path: str) -> dict:
    """使用 pdfplumber 提取 PDF 文本（备选方案）"""
    try:
        import pdfplumber

        text_parts = []
        page_count = 0

        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            for page_num, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text:
                    text_parts.append(f"[第{page_num + 1}页]\n{text}")

        return {
            "text": "\n\n".join(text_parts),
            "metadata": {
                "page_count": page_count,
                "extractor": "pdfplumber",
                "char_count": sum(len(t) for t in text_parts)
            },
            "success": True
        }
    except ImportError:
        return {"success": False, "error": "pdfplumber 未安装，请运行: pip install pdfplumber"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def extract_docx(file_path: str) -> dict:
    """使用 python-docx 提取 Word 文档文本"""
    try:
        from docx import Document

        doc = Document(file_path)
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        return {
            "text": "\n".join(text_parts),
            "metadata": {
                "extractor": "python-docx",
                "char_count": sum(len(t) for t in text_parts),
                "para_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            },
            "success": True
        }
    except ImportError:
        return {"success": False, "error": "python-docx 未安装，请运行: pip install python-docx"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def extract_doc(file_path: str) -> dict:
    """提取旧版 .doc 文件文本"""
    # 尝试使用 antiword (Linux/Mac)
    import subprocess

    try:
        result = subprocess.run(
            ["antiword", file_path],
            capture_output=True,
            text=True,
            timeout=30
        )
        if result.returncode == 0:
            return {
                "text": result.stdout,
                "metadata": {"extractor": "antiword"},
                "success": True
            }
    except FileNotFoundError:
        pass
    except Exception:
        pass

    # 尝试使用 python-docx (某些 .doc 实际是 .docx)
    try:
        return extract_docx(file_path)
    except:
        pass

    return {
        "success": False,
        "error": "无法提取 .doc 文件，请安装 antiword 或转换为 .docx"
    }


def extract_text(file_path: str) -> dict:
    """根据文件类型选择提取方法"""
    ext = Path(file_path).suffix.lower()

    extractors = {
        ".pdf": [extract_pdf_pymupdf, extract_pdf_pdfplumber],
        ".docx": [extract_docx],
        ".doc": [extract_doc],
    }

    if ext not in extractors:
        return {
            "success": False,
            "error": f"不支持的文件类型: {ext}"
        }

    # 尝试所有可用的提取器
    for extractor in extractors[ext]:
        result = extractor(file_path)
        if result.get("success"):
            return result

    # 所有提取器都失败
    return {
        "success": False,
        "error": "所有提取方法均失败"
    }


def main():
    if len(sys.argv) < 2:
        print("用法: python extract_text.py <file_path> [--format json|text]")
        sys.exit(1)

    file_path = sys.argv[1]
    output_format = "json"

    if "--format" in sys.argv:
        idx = sys.argv.index("--format")
        if idx + 1 < len(sys.argv):
            output_format = sys.argv[idx + 1]

    if not os.path.exists(file_path):
        result = {"success": False, "error": f"文件不存在: {file_path}"}
    else:
        result = extract_text(file_path)

    if output_format == "json":
        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        if result.get("success"):
            print(result["text"])
        else:
            print(f"错误: {result.get('error')}", file=sys.stderr)
            sys.exit(1)


if __name__ == "__main__":
    main()
